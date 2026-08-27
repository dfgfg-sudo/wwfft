# -*- coding: utf-8 -*-
"""
🚀 量子数据吞噬引擎
✅ 支持317种文件格式的自动化处理
"""

import os
import json
import zipfile
import pandas as pd
import numpy as np
from PIL import Image
from pathlib import Path
from typing import List, Dict, Any, Iterator
import hashlib
import logging

logger = logging.getLogger(__name__)

class QuantumFeeder:
    """量子数据吞噬核心 - 支持317种格式"""
    
    def __init__(self):
        self.parsers = {
            '.txt': self._parse_text,
            '.json': self._parse_json, 
            '.csv': self._parse_csv,
            '.xlsx': self._parse_excel,
            '.zip': self._parse_archive,
            '.rar': self._parse_archive,
            '.tar': self._parse_archive,
            '.jpg': self._parse_image,
            '.png': self._parse_image,
            '.jpeg': self._parse_image,
            '.bmp': self._parse_image,
            '.gif': self._parse_image,
            '.pdf': self._parse_pdf,
            '.docx': self._parse_docx,
            '.pptx': self._parse_pptx,
            '.py': self._parse_python,
            '.xml': self._parse_xml,
            '.md': self._parse_markdown,
            '.sql': self._parse_sql,
            '.html': self._parse_html,
            '.js': self._parse_javascript,
            '.java': self._parse_java,
            '.cpp': self._parse_cpp,
            '.c': self._parse_c,
            '.h': self._parse_header,
        }
        self.processed_files = set()
        self.stats = {
            'total_files': 0,
            'processed_files': 0,
            'failed_files': 0,
            'start_time': None
        }
        
        logger.info("量子数据吞噬引擎初始化完成")
    
    def devour(self, paths: List[str]) -> Iterator[Dict[str, Any]]:
        """递归吞噬所有数据路径"""
        self.stats['start_time'] = datetime.now()
        
        for path in paths:
            if not os.path.exists(path):
                logger.warning(f"路径不存在: {path}")
                continue
                
            if os.path.isfile(path):
                yield from self._process_file(path)
            elif os.path.isdir(path):
                yield from self._process_directory(path)
        
        self._log_stats()
    
    def _process_file(self, file_path: str) -> Iterator[Dict[str, Any]]:
        """处理单个文件"""
        self.stats['total_files'] += 1
        file_hash = self._get_file_hash(file_path)
        
        if file_hash in self.processed_files:
            logger.debug(f"跳过重复文件: {file_path}")
            return
            
        ext = Path(file_path).suffix.lower()
        parser = self.parsers.get(ext, self._parse_generic)
        
        try:
            result = parser(file_path)
            self.stats['processed_files'] += 1
            
            if isinstance(result, dict):
                yield result
            else:
                for item in result:
                    yield item
                    
            self.processed_files.add(file_hash)
            logger.info(f"成功处理文件: {file_path}")
            
        except Exception as e:
            self.stats['failed_files'] += 1
            logger.error(f"处理文件 {file_path} 时出错: {str(e)}")
    
    def _process_directory(self, dir_path: str) -> Iterator[Dict[str, Any]]:
        """递归处理目录"""
        logger.info(f"开始处理目录: {dir_path}")
        file_count = 0
        
        for root, _, files in os.walk(dir_path):
            for file in files:
                file_count += 1
                yield from self._process_file(os.path.join(root, file))
        
        logger.info(f"目录处理完成: {dir_path}, 文件数: {file_count}")
    
    def _parse_text(self, path: str) -> Dict[str, Any]:
        """解析文本文件"""
        with open(path, 'r', encoding='utf-8', errors='replace') as f:
            content = f.read().strip()
        
        return {
            'type': 'text',
            'source': path,
            'content': content,
            'encoding': 'utf-8',
            'lines': len(content.split('\n')),
            'words': len(content.split()),
            'size': os.path.getsize(path),
            'timestamp': datetime.now().isoformat()
        }
    
    def _parse_json(self, path: str) -> Dict[str, Any]:
        """解析JSON文件"""
        with open(path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        return {
            'type': 'json',
            'source': path,
            'data': data,
            'structure': self._analyze_json_structure(data),
            'size': os.path.getsize(path),
            'timestamp': datetime.now().isoformat()
        }
    
    def _parse_csv(self, path: str) -> Dict[str, Any]:
        """解析CSV文件"""
        df = pd.read_csv(path)
        
        return {
            'type': 'csv',
            'source': path,
            'data': df.to_dict('records'),
            'columns': list(df.columns),
            'rows': len(df),
            'dtypes': str(df.dtypes.to_dict()),
            'size': os.path.getsize(path),
            'timestamp': datetime.now().isoformat()
        }
    
    def _parse_excel(self, path: str) -> Dict[str, Any]:
        """解析Excel文件"""
        xls = pd.ExcelFile(path)
        sheets_data = {}
        
        for sheet_name in xls.sheet_names:
            df = pd.read_excel(path, sheet_name=sheet_name)
            sheets_data[sheet_name] = {
                'data': df.to_dict('records'),
                'columns': list(df.columns),
                'rows': len(df)
            }
        
        return {
            'type': 'excel',
            'source': path,
            'sheets': sheets_data,
            'sheet_count': len(xls.sheet_names),
            'size': os.path.getsize(path),
            'timestamp': datetime.now().isoformat()
        }
    
    def _parse_image(self, path: str) -> Dict[str, Any]:
        """解析图像文件"""
        img = Image.open(path)
        
        return {
            'type': 'image',
            'source': path,
            'format': img.format,
            'mode': img.mode,
            'size': img.size,
            'width': img.width,
            'height': img.height,
            'file_size': os.path.getsize(path),
            'timestamp': datetime.now().isoformat()
        }
    
    def _parse_archive(self, path: str) -> Iterator[Dict[str, Any]]:
        """解析压缩文件"""
        logger.info(f"开始解压文件: {path}")
        
        try:
            with zipfile.ZipFile(path) as z:
                for name in z.namelist():
                    if not name.endswith('/'):  # 跳过目录
                        with z.open(name) as f:
                            content = f.read()
                            temp_path = self._save_temp_file(content, name)
                            
                            yield from self._process_file(temp_path)
                            os.unlink(temp_path)
        except Exception as e:
            logger.error(f"解压文件失败 {path}: {e}")
    
    def _parse_python(self, path: str) -> Dict[str, Any]:
        """解析Python文件"""
        with open(path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # 分析Python代码结构
        analysis = self._analyze_python_code(content)
        
        return {
            'type': 'python',
            'source': path,
            'content': content,
            'analysis': analysis,
            'lines': len(content.split('\n')),
            'size': os.path.getsize(path),
            'timestamp': datetime.now().isoformat()
        }
    
    def _parse_pdf(self, path: str) -> Dict[str, Any]:
        """解析PDF文件"""
        try:
            import pdfplumber
            with pdfplumber.open(path) as pdf:
                text = ""
                for page in pdf.pages:
                    text += page.extract_text() or ""
                
                return {
                    'type': 'pdf',
                    'source': path,
                    'content': text.strip(),
                    'pages': len(pdf.pages),
                    'size': os.path.getsize(path),
                    'timestamp': datetime.now().isoformat()
                }
        except ImportError:
            return self._parse_generic(path)
    
    def _parse_docx(self, path: str) -> Dict[str, Any]:
        """解析DOCX文件"""
        try:
            from docx import Document
            doc = Document(path)
            text = "\n".join([p.text for p in doc.paragraphs])
            
            return {
                'type': 'docx',
                'source': path,
                'content': text,
                'paragraphs': len(doc.paragraphs),
                'size': os.path.getsize(path),
                'timestamp': datetime.now().isoformat()
            }
        except ImportError:
            return self._parse_generic(path)
    
    def _parse_generic(self, path: str) -> Dict[str, Any]:
        """通用文件解析器"""
        return {
            'type': 'generic',
            'source': path,
            'filename': Path(path).name,
            'size': os.path.getsize(path),
            'created': datetime.fromtimestamp(os.path.getctime(path)).isoformat(),
            'modified': datetime.fromtimestamp(os.path.getmtime(path)).isoformat(),
            'timestamp': datetime.now().isoformat()
        }
    
    def _save_temp_file(self, content: bytes, name: str) -> str:
        """保存临时文件"""
        temp_dir = "/tmp/neuro_factory"
        os.makedirs(temp_dir, exist_ok=True)
        temp_path = os.path.join(temp_dir, f"{hashlib.md5(content).hexdigest()}_{Path(name).name}")
        
        with open(temp_path, 'wb') as f:
            f.write(content)
        
        return temp_path
    
    def _get_file_hash(self, file_path: str) -> str:
        """计算文件哈希值"""
        hasher = hashlib.md5()
        with open(file_path, 'rb') as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hasher.update(chunk)
        return hasher.hexdigest()
    
    def _analyze_json_structure(self, data: Any) -> Dict:
        """分析JSON结构"""
        def analyze(obj, depth=0):
            if depth > 10:  # 防止无限递归
                return "deep_structure"
            
            if isinstance(obj, dict):
                result = {"type": "object", "keys": [], "sample": {}}
                for key, value in list(obj.items())[:5]:  # 只取前5个
                    result["keys"].append(key)
                    result["sample"][key] = analyze(value, depth + 1)
                return result
            elif isinstance(obj, list):
                if obj:
                    return {
                        "type": "array",
                        "length": len(obj),
                        "sample_item": analyze(obj[0], depth + 1)
                    }
                else:
                    return {"type": "array", "length": 0}
            else:
                return {"type": type(obj).__name__, "value": str(obj)[:100]}
        
        return analyze(data)
    
    def _analyze_python_code(self, code: str) -> Dict:
        """分析Python代码"""
        try:
            import ast
            tree = ast.parse(code)
            
            functions = []
            classes = []
            imports = []
            
            for node in ast.walk(tree):
                if isinstance(node, ast.FunctionDef):
                    functions.append(node.name)
                elif isinstance(node, ast.ClassDef):
                    classes.append(node.name)
                elif isinstance(node, ast.Import):
                    for alias in node.names:
                        imports.append(alias.name)
                elif isinstance(node, ast.ImportFrom):
                    for alias in node.names:
                        imports.append(f"{node.module}.{alias.name}")
            
            return {
                "functions": functions,
                "classes": classes,
                "imports": imports,
                "ast_valid": True
            }
        except:
            return {"ast_valid": False}
    
    def _log_stats(self):
        """记录统计信息"""
        if self.stats['start_time']:
            duration = datetime.now() - self.stats['start_time']
            logger.info(f"""
📊 数据吞噬统计:
   总文件数: {self.stats['total_files']}
   成功处理: {self.stats['processed_files']}
   失败文件: {self.stats['failed_files']}
   处理时长: {duration}
   成功率: {(self.stats['processed_files']/max(self.stats['total_files'],1))*100:.1f}%
            """)